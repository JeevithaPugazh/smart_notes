from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document


def _timestamped_path(out_dir: Path, stem: str, suffix: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return out_dir / f"{stem}_{ts}{suffix}"


def _md_bold_to_html(text: str) -> str:
    """Convert **bold** markers to ReportLab <b> tags."""
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


def _add_runs_with_bold(para, text: str) -> None:
    """Split text on **bold** markers and add runs with bold formatting to a DOCX paragraph."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def generate_pdf_from_text(text: str, out_dir: Path, stem: str = "smart_notes") -> Path:
    """Generate a styled PDF that renders Markdown headings and bold text."""
    out_path = _timestamped_path(out_dir, stem, ".pdf")

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    base = getSampleStyleSheet()

    h1 = ParagraphStyle("H1", parent=base["Heading1"], fontSize=22, spaceAfter=10, spaceBefore=14)
    h2 = ParagraphStyle("H2", parent=base["Heading2"], fontSize=16, spaceAfter=8, spaceBefore=10)
    h3 = ParagraphStyle("H3", parent=base["Heading3"], fontSize=13, spaceAfter=6, spaceBefore=8)
    body = ParagraphStyle("Body", parent=base["Normal"], fontSize=11, leading=17, spaceAfter=5)
    bullet = ParagraphStyle(
        "Bullet", parent=base["Normal"], fontSize=11, leading=17,
        spaceAfter=4, leftIndent=20, bulletIndent=8,
    )

    story = []

    for line in (text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue

        if stripped.startswith("### "):
            story.append(Paragraph(_md_bold_to_html(stripped[4:]), h3))
        elif stripped.startswith("## "):
            story.append(Paragraph(_md_bold_to_html(stripped[3:]), h2))
        elif stripped.startswith("# "):
            story.append(Paragraph(_md_bold_to_html(stripped[2:]), h1))
        elif re.match(r"^[-*•] ", stripped):
            content = _md_bold_to_html(stripped[2:].lstrip())
            story.append(Paragraph(f"• {content}", bullet))
        else:
            story.append(Paragraph(_md_bold_to_html(stripped), body))

    doc.build(story)
    return out_path


def generate_docx_from_text(text: str, out_dir: Path, stem: str = "smart_notes") -> Path:
    """Generate a styled DOCX that renders Markdown headings and bold text."""
    out_path = _timestamped_path(out_dir, stem, ".docx")

    doc = Document()

    for line in (text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^[-*•] ", stripped):
            para = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(para, stripped[2:].lstrip())
        else:
            para = doc.add_paragraph(style="Normal")
            _add_runs_with_bold(para, stripped)

    doc.save(out_path)
    return out_path
